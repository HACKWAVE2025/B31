"""
Document Processing Service
Handles PDF, DOCX, and other document parsing
"""

import PyPDF2
import docx2txt
import pdfplumber
from docx import Document
import os
import tempfile
from PIL import Image

# pytesseract is optional for OCR
try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False
    print("Warning: pytesseract not available. OCR features will be disabled.")

class DocumentProcessor:
    
    @staticmethod
    def extract_text_from_pdf(file_path):
        """Extract text from PDF using PyPDF2 and pdfplumber"""
        text = ""
        images = []
        
        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    
                    # Extract images
                    if hasattr(page, 'images'):
                        for img in page.images:
                            images.append({
                                'page': page_num + 1,
                                'bbox': img.get('bbox', None)
                            })
        except Exception as e:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
            except Exception as pdf_error:
                raise ValueError(f"Error extracting PDF: {str(pdf_error)}")
        
        return {
            'text': text.strip(),
            'page_count': len(text.split('--- Page')),
            'images': images,
            'has_images': len(images) > 0
        }
    
    @staticmethod
    def extract_text_from_docx(file_path):
        """Extract text from DOCX"""
        try:
            # Use docx2txt for simple extraction
            text = docx2txt.extract(file_path)
            
            # Use python-docx for more detailed extraction
            doc = Document(file_path)
            
            # Extract images
            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images.append({
                        'type': 'embedded',
                        'target': rel.target_ref
                    })
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                'text': text.strip(),
                'paragraph_count': len(doc.paragraphs),
                'tables': tables,
                'images': images,
                'has_images': len(images) > 0,
                'has_tables': len(tables) > 0
            }
        except Exception as e:
            raise ValueError(f"Error extracting DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_path):
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                'text': text.strip(),
                'line_count': len(text.split('\n')),
                'word_count': len(text.split())
            }
        except Exception as e:
            raise ValueError(f"Error reading TXT: {str(e)}")
    
    @staticmethod
    def extract_text_from_image(file_path):
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            
            if not PYTESSERACT_AVAILABLE:
                return {
                    'text': '',
                    'image_size': image.size,
                    'image_format': image.format,
                    'ocr_used': False,
                    'error': 'pytesseract not available - install it for OCR support'
                }
            
            text = pytesseract.image_to_string(image)
            
            return {
                'text': text.strip(),
                'image_size': image.size,
                'image_format': image.format,
                'ocr_used': True
            }
        except Exception as e:
            raise ValueError(f"Error performing OCR: {str(e)}")
    
    @staticmethod
    def process_document(file_path, file_type):
        """Main document processing method"""
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_type.lower()
        
        if file_extension == 'pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_extension in ['docx', 'doc']:
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif file_extension == 'txt':
            return DocumentProcessor.extract_text_from_txt(file_path)
        elif file_extension in ['png', 'jpg', 'jpeg', 'gif']:
            return DocumentProcessor.extract_text_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    @staticmethod
    def get_document_metadata(file_path):
        """Get document metadata"""
        try:
            stats = os.stat(file_path)
            return {
                'file_size': stats.st_size,
                'file_size_mb': round(stats.st_size / (1024 * 1024), 2),
                'created_at': stats.st_ctime,
                'modified_at': stats.st_mtime
            }
        except Exception as e:
            return {}
